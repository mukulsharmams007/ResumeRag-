from flask import Flask, request, jsonify, send_from_directory, session, send_file
from flask_cors import CORS
import os
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
import re
import PyPDF2
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.schema import Document
import json
from datetime import datetime
import sqlite3
import pandas as pd

app = Flask(__name__)
app.secret_key = 'your-secret-key-change-this-in-production'
CORS(app, supports_credentials=True)

# Configuration
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
DB_PATH = os.path.join(BASE_DIR, 'database.db')
VECTORDB_PATH = os.path.join(BASE_DIR, 'vectordb')
EXCEL_EXPORT_PATH = os.path.join(BASE_DIR, 'excel_exports')
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt', 'doc'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

# Create all necessary folders
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(os.path.join(VECTORDB_PATH, 'resumes'), exist_ok=True)
os.makedirs(os.path.join(VECTORDB_PATH, 'jobs'), exist_ok=True)
os.makedirs(EXCEL_EXPORT_PATH, exist_ok=True)

# Initialize Database
def init_db():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    
    c.execute('''CREATE TABLE IF NOT EXISTS users
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  name TEXT NOT NULL,
                  email TEXT UNIQUE NOT NULL,
                  password TEXT NOT NULL,
                  created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)''')
    
    c.execute('''CREATE TABLE IF NOT EXISTS resumes
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  filename TEXT NOT NULL,
                  name TEXT,
                  email TEXT,
                  phone TEXT,
                  college TEXT,
                  degree TEXT,
                  skills TEXT,
                  experience TEXT,
                  education TEXT,
                  raw_text TEXT,
                  uploaded_by TEXT,
                  uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)''')
    
    c.execute('''CREATE TABLE IF NOT EXISTS jobs
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  title TEXT NOT NULL,
                  company TEXT NOT NULL,
                  location TEXT NOT NULL,
                  description TEXT,
                  requirements TEXT,
                  posted_by TEXT,
                  posted_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)''')
    
    c.execute('''CREATE TABLE IF NOT EXISTS college_students
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  name TEXT NOT NULL,
                  email TEXT NOT NULL,
                  college TEXT NOT NULL,
                  degree TEXT NOT NULL,
                  year TEXT NOT NULL,
                  phone TEXT NOT NULL,
                  created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)''')
    
    c.execute('''CREATE TABLE IF NOT EXISTS admin_contacts
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  name TEXT NOT NULL,
                  email TEXT NOT NULL,
                  phone TEXT NOT NULL,
                  message TEXT,
                  created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                  status TEXT DEFAULT 'pending')''')
    
    conn.commit()
    conn.close()

init_db()

def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def save_user(name, email, password):
    conn = get_db()
    c = conn.cursor()
    c.execute('INSERT INTO users (name, email, password) VALUES (?, ?, ?)',
              (name, email, generate_password_hash(password)))
    conn.commit()
    conn.close()

def get_user(email):
    conn = get_db()
    c = conn.cursor()
    c.execute('SELECT * FROM users WHERE email = ?', (email,))
    user = c.fetchone()
    conn.close()
    return dict(user) if user else None

def save_resume(resume_data, filename, user_email, college="", degree=""):
    conn = get_db()
    c = conn.cursor()
    c.execute('''INSERT INTO resumes 
                 (filename, name, email, phone, college, degree, skills, experience, education, raw_text, uploaded_by)
                 VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''',
              (filename, resume_data.get('name'), resume_data.get('email'), 
               resume_data.get('phone'), college, degree, json.dumps(resume_data.get('skills')),
               resume_data.get('experience'), resume_data.get('education'),
               resume_data.get('raw_text'), user_email))
    conn.commit()
    conn.close()
    export_to_excel()

def save_job(job_data, user_email):
    conn = get_db()
    c = conn.cursor()
    c.execute('''INSERT INTO jobs (title, company, location, description, requirements, posted_by)
                 VALUES (?, ?, ?, ?, ?, ?)''',
              (job_data['title'], job_data['company'], job_data['location'],
               job_data['description'], job_data['requirements'], user_email))
    conn.commit()
    conn.close()
    export_to_excel()

def get_all_resumes():
    conn = get_db()
    c = conn.cursor()
    c.execute('SELECT * FROM resumes ORDER BY uploaded_at DESC')
    resumes = [dict(row) for row in c.fetchall()]
    conn.close()
    return resumes

def get_all_jobs():
    conn = get_db()
    c = conn.cursor()
    c.execute('SELECT * FROM jobs ORDER BY posted_at DESC')
    jobs = [dict(row) for row in c.fetchall()]
    conn.close()
    return jobs

def export_to_excel():
    """Export all database data to ONE Excel file with multiple sheets"""
    try:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        conn = get_db()
        
        # Create Excel writer
        excel_file = os.path.join(EXCEL_EXPORT_PATH, f'ResumeRAG_Data_{timestamp}.xlsx')
        latest_file = os.path.join(EXCEL_EXPORT_PATH, 'ResumeRAG_Data_Latest.xlsx')
        
        with pd.ExcelWriter(excel_file, engine='openpyxl') as writer:
            # Users Sheet
            try:
                users_df = pd.read_sql_query("SELECT id, name, email, created_at FROM users", conn)
                if not users_df.empty:
                    users_df.to_excel(writer, sheet_name='Users', index=False)
            except Exception as e:
                print(f"Users sheet error: {e}")
            
            # Resumes Sheet
            try:
                resumes_df = pd.read_sql_query("SELECT * FROM resumes", conn)
                if not resumes_df.empty:
                    resumes_df.to_excel(writer, sheet_name='Resumes', index=False)
            except Exception as e:
                print(f"Resumes sheet error: {e}")
            
            # Jobs Sheet
            try:
                jobs_df = pd.read_sql_query("SELECT * FROM jobs", conn)
                if not jobs_df.empty:
                    jobs_df.to_excel(writer, sheet_name='Jobs', index=False)
            except Exception as e:
                print(f"Jobs sheet error: {e}")
            
            # Students Sheet
            try:
                students_df = pd.read_sql_query("SELECT * FROM college_students", conn)
                if not students_df.empty:
                    students_df.to_excel(writer, sheet_name='Students', index=False)
            except Exception as e:
                print(f"Students sheet error: {e}")
            
            # Admin Contacts Sheet
            try:
                contacts_df = pd.read_sql_query("SELECT * FROM admin_contacts", conn)
                if not contacts_df.empty:
                    contacts_df.to_excel(writer, sheet_name='Admin_Contacts', index=False)
            except Exception as e:
                print(f"Contacts sheet error: {e}")
        
        # Copy to latest
        import shutil
        shutil.copy(excel_file, latest_file)
        
        conn.close()
        
        print(f"Excel export completed: {excel_file}")
        print(f"Latest file: {latest_file}")
        
    except Exception as e:
        print(f"Error exporting to Excel: {str(e)}")

class ResumeRAG:
    def __init__(self):
        self.embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2"
        )
        
        self.resume_db = Chroma(
            collection_name="resumes",
            embedding_function=self.embeddings,
            persist_directory=os.path.join(VECTORDB_PATH, 'resumes')
        )
        
        self.job_db = Chroma(
            collection_name="jobs",
            embedding_function=self.embeddings,
            persist_directory=os.path.join(VECTORDB_PATH, 'jobs')
        )
    
    def add_resume(self, resume_data, filename, user_email):
        text = f"""
        Name: {resume_data.get('name', '')}
        Email: {resume_data.get('email', '')}
        Phone: {resume_data.get('phone', '')}
        Skills: {', '.join(resume_data.get('skills', []))}
        Experience: {resume_data.get('experience', '')}
        Education: {resume_data.get('education', '')}
        Raw Text: {resume_data.get('raw_text', '')}
        """
        
        doc = Document(
            page_content=text,
            metadata={
                'filename': filename,
                'name': resume_data.get('name', ''),
                'email': resume_data.get('email', ''),
                'phone': resume_data.get('phone', ''),
                'skills': json.dumps(resume_data.get('skills', [])),
                'uploaded_by': user_email,
                'uploaded_at': datetime.now().isoformat(),
                'type': 'resume'
            }
        )
        
        self.resume_db.add_documents([doc])
        self.resume_db.persist()
    
    def add_job(self, job_data, user_email):
        text = f"""
        Title: {job_data.get('title', '')}
        Company: {job_data.get('company', '')}
        Location: {job_data.get('location', '')}
        Description: {job_data.get('description', '')}
        Requirements: {job_data.get('requirements', '')}
        """
        
        doc = Document(
            page_content=text,
            metadata={
                'title': job_data.get('title', ''),
                'company': job_data.get('company', ''),
                'location': job_data.get('location', ''),
                'posted_by': user_email,
                'posted_at': datetime.now().isoformat(),
                'type': 'job'
            }
        )
        
        self.job_db.add_documents([doc])
        self.job_db.persist()
    
    def search_resumes(self, job_description, top_k=5):
        results = self.resume_db.similarity_search_with_score(
            job_description,
            k=top_k
        )
        
        matches = []
        for doc, score in results:
            skills_str = doc.metadata.get('skills', '[]')
            try:
                skills = json.loads(skills_str) if skills_str else []
            except:
                skills = []
            
            matches.append({
                'filename': doc.metadata.get('filename'),
                'name': doc.metadata.get('name'),
                'email': doc.metadata.get('email'),
                'phone': doc.metadata.get('phone'),
                'skills': skills,
                'match_score': float(1 - score),
                'preview': doc.page_content[:200]
            })
        
        return matches
    
    def match_jobs(self, resume_text, top_k=5):
        results = self.job_db.similarity_search_with_score(
            resume_text,
            k=top_k
        )
        
        matches = []
        for doc, score in results:
            matches.append({
                'title': doc.metadata.get('title'),
                'company': doc.metadata.get('company'),
                'location': doc.metadata.get('location'),
                'match_score': float(1 - score),
                'description': doc.page_content[:300]
            })
        
        return matches

rag_engine = ResumeRAG()

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(filepath):
    try:
        text = ""
        with open(filepath, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text
    except Exception as e:
        print(f"Error reading PDF: {str(e)}")
        return ""

def extract_text_from_docx(filepath):
    try:
        doc = DocxDocument(filepath)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        print(f"Error reading DOCX: {str(e)}")
        return ""

def extract_text_from_txt(filepath):
    try:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as file:
            return file.read()
    except Exception as e:
        print(f"Error reading TXT: {str(e)}")
        return ""

def extract_resume_data(text, filename):
    if not text:
        return {
            'name': 'Unknown',
            'email': '',
            'phone': '',
            'skills': [],
            'experience': '',
            'education': '',
            'raw_text': 'Could not extract text from file'
        }
    
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    name = "Unknown"
    for line in lines[:5]:
        if len(line) > 3 and len(line) < 50 and not re.search(r'[@\d]', line):
            name = line
            break
    
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    email_match = re.search(email_pattern, text)
    email = email_match.group(0) if email_match else ""
    
    phone_patterns = [
        r'\+?\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
        r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
        r'\d{10}',
    ]
    phone = ""
    for pattern in phone_patterns:
        phone_match = re.search(pattern, text)
        if phone_match:
            phone = phone_match.group(0)
            break
    
    common_skills = [
        'python', 'java', 'javascript', 'c++', 'c#', 'ruby', 'php', 'swift', 'kotlin',
        'typescript', 'go', 'rust', 'scala', 'html', 'css', 'react', 'angular', 'vue',
        'node.js', 'express', 'django', 'flask', 'spring', 'asp.net', 'bootstrap',
        'tailwind', 'sql', 'mysql', 'postgresql', 'mongodb', 'oracle', 'redis',
        'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'git', 'ci/cd',
        'machine learning', 'deep learning', 'tensorflow', 'pytorch', 'pandas',
        'numpy', 'ai', 'nlp', 'agile', 'scrum', 'rest api', 'microservices'
    ]
    
    text_lower = text.lower()
    skills = []
    for skill in common_skills:
        if skill in text_lower:
            skills.append(skill.title())
    
    skills = sorted(list(set(skills)))
    
    experience = extract_section(text, ['experience', 'work history', 'employment'])
    education = extract_section(text, ['education', 'academic', 'qualification'])
    
    return {
        'name': name,
        'email': email,
        'phone': phone,
        'skills': skills,
        'experience': experience,
        'education': education,
        'raw_text': text[:500]
    }

def extract_section(text, keywords):
    text_lower = text.lower()
    lines = text.split('\n')
    
    for keyword in keywords:
        for i, line in enumerate(lines):
            if keyword in line.lower():
                section_lines = lines[i:min(i+10, len(lines))]
                section_text = '\n'.join(section_lines)
                return section_text[:300]
    
    return ""

@app.route('/')
def index():
    frontend_folder = os.path.join(BASE_DIR, 'frontend')
    return send_from_directory(frontend_folder, 'index.html')

@app.route('/<path:path>')
def serve_static(path):
    frontend_folder = os.path.join(BASE_DIR, 'frontend')
    return send_from_directory(frontend_folder, path)

@app.route('/api/signup', methods=['POST'])
def signup():
    try:
        data = request.get_json()
        email = data.get('email')
        password = data.get('password')
        name = data.get('name')
        
        if not email or not password or not name:
            return jsonify({'error': 'All fields required'}), 400
        
        if get_user(email):
            return jsonify({'error': 'User already exists'}), 400
        
        save_user(name, email, password)
        export_to_excel()
        
        session['user_email'] = email
        session['user_name'] = name
        
        return jsonify({
            'success': True,
            'message': 'Account created successfully',
            'user': {'email': email, 'name': name}
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/login', methods=['POST'])
def login():
    try:
        data = request.get_json()
        email = data.get('email')
        password = data.get('password')
        
        if not email or not password:
            return jsonify({'error': 'Email and password required'}), 400
        
        user = get_user(email)
        if not user or not check_password_hash(user['password'], password):
            return jsonify({'error': 'Invalid credentials'}), 401
        
        session['user_email'] = email
        session['user_name'] = user['name']
        
        return jsonify({
            'success': True,
            'message': 'Login successful',
            'user': {'email': email, 'name': user['name']}
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/logout', methods=['POST'])
def logout():
    session.clear()
    return jsonify({'success': True, 'message': 'Logged out'})

@app.route('/api/check-auth', methods=['GET'])
def check_auth():
    if 'user_email' in session:
        return jsonify({
            'authenticated': True,
            'user': {
                'email': session['user_email'],
                'name': session['user_name']
            }
        })
    return jsonify({'authenticated': False})

@app.route('/api/upload-resume', methods=['POST'])
def upload_resume():
    try:
        if 'user_email' not in session:
            return jsonify({'error': 'Please login first'}), 401
        
        if 'resume' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['resume']
        college = request.form.get('college', '')
        degree = request.form.get('degree', '')
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            name, ext = os.path.splitext(filename)
            unique_filename = f"{name}_{timestamp}{ext}"
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
            
            file.save(filepath)
            print(f"File saved to: {filepath}")
            
            file_extension = unique_filename.rsplit('.', 1)[1].lower()
            
            if file_extension == 'pdf':
                text = extract_text_from_pdf(filepath)
            elif file_extension in ['docx', 'doc']:
                text = extract_text_from_docx(filepath)
            else:
                text = extract_text_from_txt(filepath)
            
            resume_data = extract_resume_data(text, unique_filename)
            resume_data['college'] = college
            resume_data['degree'] = degree
            
            save_resume(resume_data, unique_filename, session['user_email'], college, degree)
            rag_engine.add_resume(resume_data, unique_filename, session['user_email'])
            
            return jsonify({
                'success': True,
                'filename': unique_filename,
                'data': resume_data
            })
        
        return jsonify({'error': 'Invalid file type'}), 400
        
    except Exception as e:
        print(f"Error: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/search-resumes', methods=['POST'])
def search_resumes():
    try:
        if 'user_email' not in session:
            return jsonify({'error': 'Please login first'}), 401
        
        data = request.get_json()
        job_description = data.get('job_description', '')
        top_k = data.get('top_k', 5)
        
        matches = rag_engine.search_resumes(job_description, top_k)
        
        return jsonify({
            'success': True,
            'matches': matches
        })
    except Exception as e:
        print(f"Error: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/add-job', methods=['POST'])
def add_job():
    try:
        if 'user_email' not in session:
            return jsonify({'error': 'Please login first'}), 401
        
        data = request.get_json()
        
        save_job(data, session['user_email'])
        rag_engine.add_job(data, session['user_email'])
        
        return jsonify({
            'success': True,
            'message': 'Job posted successfully'
        })
    except Exception as e:
        print(f"Error: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/match-jobs', methods=['POST'])
def match_jobs():
    try:
        if 'user_email' not in session:
            return jsonify({'error': 'Please login first'}), 401
        
        data = request.get_json()
        resume_text = data.get('resume_text', '')
        
        matches = rag_engine.match_jobs(resume_text, 5)
        
        return jsonify({
            'success': True,
            'jobs': matches
        })
    except Exception as e:
        print(f"Error: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/get-resumes', methods=['GET'])
def get_resumes():
    try:
        if 'user_email' not in session:
            return jsonify({'error': 'Please login first'}), 401
        
        resumes = get_all_resumes()
        for resume in resumes:
            if resume.get('skills'):
                resume['skills'] = json.loads(resume['skills'])
        
        return jsonify({
            'success': True,
            'resumes': resumes
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/get-jobs', methods=['GET'])
def get_jobs():
    try:
        if 'user_email' not in session:
            return jsonify({'error': 'Please login first'}), 401
        
        jobs = get_all_jobs()
        
        return jsonify({
            'success': True,
            'jobs': jobs
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/add-student', methods=['POST'])
def add_student():
    try:
        if 'user_email' not in session:
            return jsonify({'error': 'Please login first'}), 401
        
        data = request.get_json()
        
        conn = get_db()
        c = conn.cursor()
        c.execute('''INSERT INTO college_students (name, email, college, degree, year, phone)
                     VALUES (?, ?, ?, ?, ?, ?)''',
                  (data['name'], data['email'], data['college'], 
                   data['degree'], data['year'], data['phone']))
        conn.commit()
        conn.close()
        
        export_to_excel()
        
        return jsonify({
            'success': True,
            'message': 'Student added successfully'
        })
    except Exception as e:
        print(f"Error: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/get-students', methods=['GET'])
def get_students():
    try:
        if 'user_email' not in session:
            return jsonify({'error': 'Please login first'}), 401
        
        conn = get_db()
        c = conn.cursor()
        c.execute('SELECT * FROM college_students ORDER BY created_at DESC')
        students = [dict(row) for row in c.fetchall()]
        conn.close()
        
        return jsonify({
            'success': True,
            'students': students
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/get-college-resumes', methods=['GET'])
def get_college_resumes():
    """Get all resumes with college information"""
    try:
        if 'user_email' not in session:
            return jsonify({'error': 'Please login first'}), 401
        
        conn = get_db()
        c = conn.cursor()
        c.execute('SELECT * FROM resumes WHERE college IS NOT NULL AND college != "" ORDER BY uploaded_at DESC')
        resumes = [dict(row) for row in c.fetchall()]
        
        for resume in resumes:
            if resume.get('skills'):
                try:
                    resume['skills'] = json.loads(resume['skills'])
                except:
                    resume['skills'] = []
        
        conn.close()
        
        return jsonify({
            'success': True,
            'resumes': resumes
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/contact-admin', methods=['POST'])
def contact_admin():
    try:
        if 'user_email' not in session:
            return jsonify({'error': 'Please login first'}), 401
        
        data = request.get_json()
        
        conn = get_db()
        c = conn.cursor()
        c.execute('''INSERT INTO admin_contacts (name, email, phone, message)
                     VALUES (?, ?, ?, ?)''',
                  (data['name'], data['email'], data['phone'], data['message']))
        conn.commit()
        conn.close()
        
        export_to_excel()
        
        return jsonify({
            'success': True,
            'message': 'Your message has been sent to admin'
        })
    except Exception as e:
        print(f"Error: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/export-excel', methods=['GET'])
def export_excel_endpoint():
    try:
        if 'user_email' not in session:
            return jsonify({'error': 'Please login first'}), 401
        
        export_to_excel()
        
        latest_file = os.path.join(EXCEL_EXPORT_PATH, 'users_latest.xlsx')
        if os.path.exists(latest_file):
            return send_file(latest_file, as_attachment=True, download_name='database_export.xlsx')
        else:
            return jsonify({'error': 'No export file available'}), 404
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/list-uploaded-files', methods=['GET'])
def list_uploaded_files():
    try:
        if 'user_email' not in session:
            return jsonify({'error': 'Please login first'}), 401
        
        files = []
        if os.path.exists(UPLOAD_FOLDER):
            for filename in os.listdir(UPLOAD_FOLDER):
                filepath = os.path.join(UPLOAD_FOLDER, filename)
                if os.path.isfile(filepath):
                    files.append({
                        'filename': filename,
                        'size': os.path.getsize(filepath),
                        'modified': datetime.fromtimestamp(os.path.getmtime(filepath)).isoformat()
                    })
        
        return jsonify({
            'success': True,
            'files': files,
            'total': len(files)
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    print("=" * 60)
    print("🚀 Starting ResumeRAG Flask Server...")
    print("=" * 60)
    print(f"📁 Upload folder: {UPLOAD_FOLDER}")
    print(f"📁 Folder exists: {os.path.exists(UPLOAD_FOLDER)}")
    print(f"💾 Database: {DB_PATH}")
    print(f"💾 Database exists: {os.path.exists(DB_PATH)}")
    print(f"📊 Excel exports: {EXCEL_EXPORT_PATH}")
    print(f"🗄️ VectorDB: {VECTORDB_PATH}")
    print(f"🌐 Server will run on: http://0.0.0.0:5000")
    print("=" * 60)
    
    if os.path.exists(DB_PATH):
        print("\n📊 Exporting initial data to Excel...")
        export_to_excel()
    
    print("\n✅ Server ready! Open http://localhost:5000")
    print("=" * 60)
    app.run(debug=True, port=5000, host='0.0.0.0')